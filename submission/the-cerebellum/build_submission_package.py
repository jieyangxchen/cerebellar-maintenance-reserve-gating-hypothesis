#!/usr/bin/env python3
"""Build deterministic Word files for The Cerebellum submission package."""

from __future__ import annotations

import hashlib
import re
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
PACKAGE_DIR = Path(__file__).resolve().parent / "package"
MANUSCRIPT = ROOT / "manuscript" / "manuscript.md"
PROTOCOLS = (
    ROOT / "protocols" / "prospective-cohort.md",
    ROOT / "protocols" / "statistical-analysis-plan.md",
    ROOT / "protocols" / "early-intervention-trial.md",
)

TITLE = "A Maintenance–Reserve–Gating Framework for Modifier Effects in Hereditary Cerebellar Ataxia"
RUNNING_TITLE = "Maintenance–reserve–gating framework in hereditary ataxia"
AUTHOR = "Jieyang Chen"
AFFILIATION = "Independent Researcher, Hangzhou, China"
EMAIL = "278404704@qq.com"
ORCID = "0009-0001-9247-2085"
REPOSITORY = "https://github.com/jieyangxchen/cerebellar-maintenance-reserve-gating-hypothesis"
VERSION = "0.2.3"
PACKAGE_DATE = "21 August 2026"
EVIDENCE_CUTOFF = "19 August 2026"
FIXED_TIME = datetime(2026, 8, 21, 0, 0, 0, tzinfo=timezone.utc)

MAIN_DOCX = PACKAGE_DIR / "01_Main_Manuscript.docx"
TITLE_DOCX = PACKAGE_DIR / "02_Title_Page.docx"
COVER_DOCX = PACKAGE_DIR / "03_Cover_Letter.docx"
SUPPLEMENT_DOCX = PACKAGE_DIR / "04_Supplementary_Protocols.docx"
OUTPUTS = (MAIN_DOCX, TITLE_DOCX, COVER_DOCX, SUPPLEMENT_DOCX)


def _set_run_font(
    run,
    *,
    name: str = "Times New Roman",
    size: float = 10,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str = "000000",
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_cell_margins(cell, *, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths_inches: list[float]) -> None:
    widths = [int(round(value * 1440)) for value in widths_inches]
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_inches[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def _prevent_row_break(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:cantSplit")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _configure_document(doc: Document, *, title: str, subject: str, header_text: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1.12)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.42)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    heading_tokens = {
        "Heading 1": (12, 12, 6, False),
        "Heading 2": (11, 9, 4, False),
        "Heading 3": (10, 7, 3, True),
    }
    for name, (size, before, after, italic) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.italic = italic
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    title_style = styles["Title"]
    title_style.font.name = "Times New Roman"
    title_style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    title_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    title_style.font.size = Pt(14)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.space_after = Pt(14)
    title_ppr = title_style._element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)

    caption = styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_together = True

    for list_style in ("List Bullet", "List Bullet 2", "List Bullet 3", "List Number", "List Number 2", "List Number 3"):
        if list_style in styles:
            style = styles[list_style]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
            style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
            style.font.size = Pt(10)
            style.paragraph_format.space_after = Pt(3)
            style.paragraph_format.line_spacing = 1.15

    doc.core_properties.title = title
    doc.core_properties.subject = subject
    doc.core_properties.author = AUTHOR
    doc.core_properties.last_modified_by = AUTHOR
    doc.core_properties.created = FIXED_TIME
    doc.core_properties.modified = FIXED_TIME
    doc.core_properties.keywords = (
        "Spinocerebellar Ataxias; Cerebellum; Environmental Exposure; "
        "Disease Progression; Biomarkers; Genetic Predisposition to Disease"
    )

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run(header_text)
    _set_run_font(run, size=8, color="666666")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("Page ")
    _set_run_font(run, size=8, color="666666")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_begin, instr, fld_separate, display, fld_end))


def _add_centered(doc: Document, text: str, *, size: float = 10, bold: bool = False, italic: bool = False, after: float = 6) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(after)
    run = paragraph.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic)


def _latex_preprocess(text: str) -> str:
    value = text.strip()
    value = value.replace("\\frac", "⟦FRAC⟧")
    value = re.sub(r"\\(?:left|right)\b", "", value)
    value = value.replace("\\,", " ").replace("\\;", " ").replace("\\ ", " ")
    value = value.replace("\\qquad", "   ").replace("\\quad", "  ")
    value = value.replace("\\mid", "|")
    value = value.replace("\\{", "⦃").replace("\\}", "⦄")
    wrappers = ("text", "mathrm", "mathbf", "boldsymbol", "operatorname")
    for wrapper in wrappers:
        pattern = re.compile(rf"\\{wrapper}\{{([^{{}}]*)\}}")
        while pattern.search(value):
            value = pattern.sub(lambda match: match.group(1), value)
    replacements = {
        "\\mathcal L": "ℒ",
        "\\Lambda": "Λ",
        "\\Delta": "Δ",
        "\\lambda": "λ",
        "\\alpha": "α",
        "\\beta": "β",
        "\\delta": "δ",
        "\\gamma": "γ",
        "\\theta": "θ",
        "\\phi": "φ",
        "\\psi": "ψ",
        "\\eta": "η",
        "\\kappa": "κ",
        "\\epsilon": "ε",
        "\\pm": "±",
        "\\ldots": "…",
        "\\log": "log",
        "\\rightarrow": "→",
        "\\ge": "≥",
        "\\le": "≤",
        "\\gg": "≫",
        "\\approx": "≈",
        "\\times": "×",
        "\\partial": "∂",
        "\\sum": "∑",
        "\\inf": "inf",
        "\\exp": "exp",
        "\\max": "max",
    }
    for source, target in sorted(replacements.items(), key=lambda item: len(item[0]), reverse=True):
        value = value.replace(source, target)
    value = value.replace("\\%", "%")
    value = value.replace("<=", "≤").replace(">=", "≥")
    value = value.replace("\\\\", " ; ")
    value = re.sub(r"\\[A-Za-z]+", "", value)
    value = value.replace("⟦FRAC⟧", "\\frac")
    for source, target in (
        ("Delta", "Δ"),
        ("alpha", "α"),
        ("beta", "β"),
        ("gamma", "γ"),
        ("theta", "θ"),
    ):
        value = re.sub(rf"(?<![A-Za-z]){source}(?=_|[^A-Za-z]|$)", target, value)
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def _math_run(text: str):
    run = OxmlElement("m:r")
    run_props = OxmlElement("m:rPr")
    normal = OxmlElement("m:nor")
    run_props.append(normal)
    run.append(run_props)
    node = OxmlElement("m:t")
    node.text = text.replace("⦃", "{").replace("⦄", "}")
    run.append(node)
    return run


class _EquationParser:
    def __init__(self, source: str):
        self.source = source
        self.index = 0

    def parse(self, stop: str | None = None) -> list:
        nodes: list = []
        while self.index < len(self.source):
            char = self.source[self.index]
            if stop is not None and char == stop:
                self.index += 1
                break
            atom = self._atom()
            if not atom:
                continue
            subscript = None
            superscript = None
            while self.index < len(self.source) and self.source[self.index] in "_^":
                marker = self.source[self.index]
                self.index += 1
                script = self._script()
                if marker == "_":
                    subscript = script
                else:
                    superscript = script
            if subscript is not None or superscript is not None:
                atom = [self._scripted(atom, subscript, superscript)]
            nodes.extend(atom)
        return nodes

    def _atom(self) -> list:
        if self.source.startswith("\\frac", self.index):
            self.index += len("\\frac")
            numerator = self._script()
            denominator = self._script()
            fraction = OxmlElement("m:f")
            num = OxmlElement("m:num")
            den = OxmlElement("m:den")
            for node in numerator:
                num.append(node)
            for node in denominator:
                den.append(node)
            fraction.extend((num, den))
            return [fraction]
        char = self.source[self.index]
        if char == "{":
            self.index += 1
            return self.parse("}")
        if char == "}":
            self.index += 1
            return []
        if char.isspace():
            start = self.index
            while self.index < len(self.source) and self.source[self.index].isspace():
                self.index += 1
            return [_math_run(self.source[start:self.index])]
        if char.isalnum() or char in "αβγδΔλθφψηκεΛℒ":
            start = self.index
            while self.index < len(self.source):
                current = self.source[self.index]
                if not (current.isalnum() or current in "αβγδΔλθφψηκεΛℒ"):
                    break
                self.index += 1
            return [_math_run(self.source[start:self.index])]
        self.index += 1
        return [_math_run(char)]

    def _script(self) -> list:
        if self.index < len(self.source) and self.source[self.index] == "{":
            self.index += 1
            return self.parse("}")
        return self._atom()

    @staticmethod
    def _scripted(base: list, subscript: list | None, superscript: list | None):
        if subscript is not None and superscript is not None:
            container = OxmlElement("m:sSubSup")
            tags = (("m:e", base), ("m:sub", subscript), ("m:sup", superscript))
        elif subscript is not None:
            container = OxmlElement("m:sSub")
            tags = (("m:e", base), ("m:sub", subscript))
        else:
            container = OxmlElement("m:sSup")
            tags = (("m:e", base), ("m:sup", superscript or []))
        for tag, nodes in tags:
            branch = OxmlElement(tag)
            for node in nodes:
                branch.append(node)
            container.append(branch)
        return container


def _add_equation(doc: Document, latex: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_together = True
    math = OxmlElement("m:oMath")
    parser = _EquationParser(_latex_preprocess(latex))
    for node in parser.parse():
        math.append(node)
    paragraph._p.append(math)


def _append_inline_math(paragraph, latex: str) -> None:
    math = OxmlElement("m:oMath")
    parser = _EquationParser(_latex_preprocess(latex))
    for node in parser.parse():
        math.append(node)
    paragraph._p.append(math)


INLINE_TOKEN = re.compile(
    r"(\*\*.+?\*\*|(?<!\w)\*[^*\n]+\*(?!\w)|`[^`]+`|\$[^$\n]+\$|\[[^\]]+\]\([^)]+\))"
)


def _strip_markdown_links(text: str) -> str:
    """Return Markdown link labels while preserving citations such as ``[1]``.

    A small balanced-parenthesis scanner is used because DOI URLs commonly
    contain parentheses, which a single regular expression truncates.
    """
    output: list[str] = []
    cursor = 0
    while cursor < len(text):
        start = text.find("[", cursor)
        if start < 0:
            output.append(text[cursor:])
            break
        output.append(text[cursor:start])
        label_end = text.find("]", start + 1)
        if label_end < 0:
            output.append(text[start:])
            break
        if label_end + 1 >= len(text) or text[label_end + 1] != "(":
            output.append(text[start : label_end + 1])
            cursor = label_end + 1
            continue
        depth = 1
        position = label_end + 2
        while position < len(text) and depth:
            if text[position] == "(":
                depth += 1
            elif text[position] == ")":
                depth -= 1
            position += 1
        if depth:
            output.append(text[start:])
            break
        output.append(text[start + 1 : label_end])
        cursor = position
    return "".join(output)


def _clean_inline(text: str) -> str:
    value = _strip_markdown_links(text)
    value = value.replace("**", "").replace("`", "")
    value = re.sub(r"(?<!\w)\*([^*\n]+)\*(?!\w)", r"\1", value)
    return value


def _add_inline_runs(paragraph, text: str, *, size: float = 10, default_bold: bool = False) -> None:
    text = _strip_markdown_links(text)
    cursor = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            _set_run_font(run, size=size, bold=default_bold)
        token = match.group(0)
        if token.startswith("**"):
            _add_inline_runs(paragraph, token[2:-2], size=size, default_bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, size=size, italic=True, bold=default_bold)
        elif token.startswith("`"):
            content = token[1:-1]
            if "Specification Memorandum" in content:
                run = paragraph.add_run(content)
                _set_run_font(run, size=size, italic=True, bold=default_bold)
            else:
                _append_inline_math(paragraph, content)
        elif token.startswith("$"):
            _append_inline_math(paragraph, token[1:-1])
        else:
            label = re.match(r"\[([^\]]+)\]", token).group(1)
            run = paragraph.add_run(label)
            _set_run_font(run, size=size, bold=default_bold)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        _set_run_font(run, size=size, bold=default_bold)


def _table_widths(column_count: int) -> list[float]:
    if column_count == 2:
        return [1.65, 4.85]
    if column_count == 3:
        return [3.4, 1.55, 1.55]
    if column_count == 4:
        return [2.8, 1.23, 1.23, 1.24]
    if column_count == 6:
        return [2.2, 0.86, 0.86, 0.86, 0.86, 0.86]
    width = 6.5 / column_count
    return [width] * column_count


def _parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _add_table(doc: Document, lines: list[str]) -> None:
    rows = [_parse_table_row(line) for line in lines]
    if len(rows) < 2:
        return
    data = [rows[0], *rows[2:]]
    column_count = len(data[0])
    table = doc.add_table(rows=len(data), cols=column_count)
    table.style = "Table Grid"
    _set_table_geometry(table, _table_widths(column_count))
    _repeat_header(table.rows[0])
    for row in table.rows:
        _prevent_row_break(row)
    for row_index, values in enumerate(data):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            _add_inline_runs(paragraph, value, size=8.5, default_bold=row_index == 0)
            if row_index == 0:
                _shade_cell(cell, "EDEDED")
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def _add_image(doc: Document, source_path: Path, alt_text: str) -> None:
    image_path = source_path
    if source_path.suffix.lower() == ".svg":
        image_path = source_path.with_suffix(".png")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.3))
    drawing = run._r.xpath(".//wp:docPr")
    if drawing:
        drawing[0].set("descr", alt_text)


def _is_table_separator(line: str) -> bool:
    cells = _parse_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def _add_markdown(
    doc: Document,
    markdown: str,
    *,
    source_path: Path,
    heading_shift: int,
    page_break_before_first: bool = False,
) -> None:
    lines = markdown.splitlines()
    index = 0
    first_heading_pending = page_break_before_first
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if not stripped or stripped == "---":
            index += 1
            continue
        if stripped == "$$":
            equation_lines: list[str] = []
            index += 1
            while index < len(lines) and lines[index].strip() != "$$":
                equation_lines.append(lines[index])
                index += 1
            _add_equation(doc, " ".join(equation_lines))
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and _is_table_separator(lines[index + 1]):
            table_lines = [raw, lines[index + 1]]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            _add_table(doc, table_lines)
            continue
        image_match = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            relative = image_match.group(2)
            _add_image(doc, (source_path.parent / relative).resolve(), image_match.group(1))
            index += 1
            continue
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            source_level = len(heading_match.group(1))
            level = min(3, max(1, source_level + heading_shift))
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            if first_heading_pending:
                paragraph.paragraph_format.page_break_before = True
                first_heading_pending = False
            _add_inline_runs(paragraph, heading_match.group(2), size={1: 12, 2: 11, 3: 10}[level], default_bold=True)
            index += 1
            continue
        if stripped.startswith(("**Figure ", "**Fig. ")):
            paragraph = doc.add_paragraph(style="Caption")
            _add_inline_runs(paragraph, stripped, size=9)
            index += 1
            continue
        list_match = re.match(r"^(\s*)[-*]\s+(.+)$", raw)
        numbered_match = re.match(r"^(\s*)(\d+)[.)]\s+(.+)$", raw)
        if list_match or numbered_match:
            match = list_match or numbered_match
            level = min(3, len(match.group(1)) // 2 + 1)
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25 * level)
            paragraph.paragraph_format.first_line_indent = Inches(-0.18)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.15
            if list_match:
                prefix = "• "
                content = list_match.group(2)
            else:
                prefix = f"{numbered_match.group(2)}. "
                content = numbered_match.group(3)
            prefix_run = paragraph.add_run(prefix)
            _set_run_font(prefix_run, size=10)
            _add_inline_runs(paragraph, content, size=10)
            index += 1
            continue
        if stripped.startswith("> "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.3)
            paragraph.paragraph_format.right_indent = Inches(0.2)
            paragraph.paragraph_format.space_after = Pt(6)
            _add_inline_runs(paragraph, stripped[2:], size=10)
            for run in paragraph.runs:
                run.italic = True
            index += 1
            continue
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.keep_together = False
        _add_inline_runs(paragraph, stripped, size=10)
        index += 1


def _extract_main_sections(text: str) -> tuple[str, str, str]:
    abstract_marker = "## Abstract"
    body_marker = "## One-sentence thesis"
    abstract_start = text.index(abstract_marker) + len(abstract_marker)
    body_start = text.index(body_marker)
    abstract_and_keywords = text[abstract_start:body_start].strip()
    keyword_match = re.search(r"\*\*Keywords:\*\*\s*(.+)$", abstract_and_keywords, re.MULTILINE)
    if keyword_match is None:
        raise ValueError("Keywords not found in manuscript")
    keywords = keyword_match.group(1).strip()
    abstract = abstract_and_keywords[:keyword_match.start()].strip()
    body = text[body_start:].strip()
    return abstract, keywords, body


def _add_main_title_page(doc: Document) -> None:
    paragraph = doc.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_inline_runs(paragraph, TITLE, size=14, default_bold=True)
    _add_centered(doc, AUTHOR, size=11, bold=True, after=4)
    _add_centered(doc, f"ORCID: {ORCID}", size=10, after=2)
    _add_centered(doc, AFFILIATION, size=10, after=2)
    _add_centered(doc, f"Correspondence: {EMAIL}", size=10, after=10)
    _add_centered(doc, f"Running title: {RUNNING_TITLE}", size=9.5, italic=True, after=3)
    _add_centered(doc, f"Manuscript version {VERSION}; evidence cut-off {EVIDENCE_CUTOFF}", size=9.5, after=3)
    _add_centered(doc, f"Public repository/preprint package: {REPOSITORY}", size=9.5, after=10)
    _add_centered(doc, "This theoretical article reports no original human-participant or animal research.", size=9.5, italic=True, after=8)
    doc.add_page_break()


def build_main_manuscript() -> Document:
    source = MANUSCRIPT.read_text(encoding="utf-8")
    abstract, keywords, body = _extract_main_sections(source)
    doc = Document()
    _configure_document(doc, title=TITLE, subject="New ideas, opinion and controversies", header_text=RUNNING_TITLE)
    _add_main_title_page(doc)
    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run("Abstract")
    for paragraph_text in re.split(r"\n\s*\n", abstract):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_inline_runs(paragraph, paragraph_text.replace("\n", " "), size=10)
    keyword_paragraph = doc.add_paragraph()
    label = keyword_paragraph.add_run("Keywords: ")
    _set_run_font(label, size=10, bold=True)
    _add_inline_runs(keyword_paragraph, keywords, size=10)
    _add_markdown(doc, body, source_path=MANUSCRIPT, heading_shift=-1)
    return doc


def build_title_page() -> Document:
    doc = Document()
    _configure_document(doc, title=f"Title page - {TITLE}", subject="Separate submission title page", header_text="Title page")
    paragraph = doc.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_inline_runs(paragraph, TITLE, size=14, default_bold=True)
    _add_centered(doc, AUTHOR, size=11, bold=True, after=3)
    _add_centered(doc, f"ORCID: {ORCID}", size=10, after=2)
    _add_centered(doc, AFFILIATION, size=10, after=2)
    _add_centered(doc, f"Corresponding author: {AUTHOR} | {EMAIL}", size=10, after=12)

    entries = (
        ("Running title", RUNNING_TITLE),
        ("Proposed article type", "New ideas, opinion and controversies"),
        ("Research status", "Theoretical framework and critical review; no original participant-level or animal research"),
        ("Manuscript version", VERSION),
        ("Evidence cut-off", EVIDENCE_CUTOFF),
        ("Public repository", REPOSITORY),
        ("Funding", "No external funding was received for this work."),
        (
            "Competing interests",
            "The author declares no competing interests. The author retains copyright in the associated public repository and may consider future commercial-licensing requests. No commercial funding or payment was received for this work.",
        ),
        (
            "Author contribution",
            "Jieyang Chen conceived the framework, performed the literature organization and claim audit, designed the proposed research programme, prepared the figures and reproducibility materials, and drafted and revised the manuscript.",
        ),
        ("Acknowledgements", "None."),
        (
            "AI-assisted tools",
            "OpenAI Codex assisted with literature organization, drafting and language editing, figure and repository-check code, and internal consistency review. It is not an author; the named author remains responsible for the submitted work.",
        ),
    )
    for label_text, value in entries:
        paragraph = doc.add_paragraph()
        label = paragraph.add_run(f"{label_text}: ")
        _set_run_font(label, size=10, bold=True)
        _add_inline_runs(paragraph, value, size=10)
    return doc


def build_cover_letter() -> Document:
    doc = Document()
    _configure_document(doc, title=f"Cover letter - {TITLE}", subject="Cover letter to The Cerebellum", header_text="Cover letter | The Cerebellum")

    for line in (PACKAGE_DATE, "Editor-in-Chief and Editors", "The Cerebellum"):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        _add_inline_runs(paragraph, line, size=10)
    doc.add_paragraph()
    salutation = doc.add_paragraph()
    _add_inline_runs(salutation, "Dear Professor Manto and Editors,", size=10)

    paragraphs = (
        f"I am submitting the manuscript entitled “{TITLE}” for consideration as a “new ideas, opinion and controversies” article. Following a presubmission inquiry, Springer Nature Publishing Support advised formal submission so that the Editor-in-Chief can assess suitability during initial screening; the support response did not select an article category, and I would welcome editorial reclassification if appropriate.",
        "The manuscript proposes a nested and falsifiable framework for timing heterogeneity in hereditary cerebellar ataxias. Its observable layer tests a prospectively frozen measured exposure score for non-linear, genotype-specific associations with phenoconversion and biomarker trajectories. Its mechanistic extension involving an unknown input, activation/exchange gating, and dynamic reserve is explicitly identified as conjecture and requires independent mediator and perturbation evidence. A statistical hump alone is not presented as evidence for gating.",
        "The work builds directly on established structural and functional concepts of cerebellar reserve and states its narrower novelty as the activation-supply mismatch prediction, the observation model linking dynamic reserve to phenoconversion, and layer-specific rejection rules. SCA3 is proposed as the primary test bed and SCA6 as a stringent transport and heterogeneity test. Known-target early-intervention examples are separated from validation of the unknown-input mechanism and are not treatment recommendations.",
        "The article reports no participant-level data and no new human or animal experiments. It contains three original, programmatically generated figures. Detailed cohort, statistical-analysis, and early-intervention protocol concepts are supplied as supporting research-planning material and are explicitly not clinical protocols ready for implementation.",
        f"The manuscript and supporting package have been publicly available in a CC BY-NC 4.0 GitHub repository ({REPOSITORY}). This history is disclosed transparently for assessment under the journal's prior-publication and preprint policies. The Publishing Support response did not make a journal-specific policy determination, so no prior approval is claimed.",
        "The manuscript is not under consideration by another journal. No acknowledgements are required. No external funding was received. The author declares no competing interests; the repository copyright and possibility of considering future commercial-licensing requests are disclosed in the manuscript. AI-assisted use of OpenAI Codex is also disclosed, and the named author retains full responsibility for source verification, scientific claims, and the submitted version.",
        "Thank you for considering the manuscript.",
    )
    for text in paragraphs:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_inline_runs(paragraph, text, size=10)
    signature = doc.add_paragraph()
    signature.paragraph_format.space_before = Pt(5)
    signature.paragraph_format.space_after = Pt(0)
    signature.paragraph_format.keep_together = True
    for index, line in enumerate(("Sincerely,", AUTHOR, AFFILIATION, EMAIL, f"ORCID: {ORCID}")):
        if index:
            signature.add_run().add_break()
        _add_inline_runs(signature, line, size=10)
    return doc


def build_supplement() -> Document:
    doc = Document()
    _configure_document(
        doc,
        title=f"Supplementary Material 1 - {TITLE}",
        subject="Research-planning protocols and SAP",
        header_text="Supplementary Material 1 | Maintenance–reserve–gating framework",
    )
    paragraph = doc.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_inline_runs(paragraph, "Supplementary Material 1: Methods and Protocol Concepts", size=14, default_bold=True)
    _add_centered(doc, TITLE, size=11, italic=True, after=4)
    _add_centered(doc, f"Version {VERSION} | Evidence cut-off {EVIDENCE_CUTOFF}", size=9.5, after=10)

    notice = doc.add_paragraph()
    notice.paragraph_format.left_indent = Inches(0.25)
    notice.paragraph_format.right_indent = Inches(0.25)
    notice.paragraph_format.space_before = Pt(6)
    notice.paragraph_format.space_after = Pt(10)
    run = notice.add_run(
        "Safety and status notice. These documents are prospective research-planning concepts. They have no sponsor authorization, regulatory approval, ethics approval, trial registration, investigational-product specification, or clinical-use authority. They do not justify self-treatment, exposure escalation, stimulation, medication, gene manipulation, or neural injury."
    )
    _set_run_font(run, size=9.5, bold=True, color="9C3F50")
    _add_centered(doc, f"Canonical public sources: {REPOSITORY}", size=9, after=4)

    for protocol in PROTOCOLS:
        markdown = protocol.read_text(encoding="utf-8")
        _add_markdown(
            doc,
            markdown,
            source_path=protocol,
            heading_shift=0,
            page_break_before_first=True,
        )
    return doc


def _deterministic_save(doc: Document, destination: Path) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="cerebellum-docx-") as temp_dir:
        raw = Path(temp_dir) / "raw.docx"
        doc.save(raw)
        with zipfile.ZipFile(raw, "r") as source, zipfile.ZipFile(
            destination,
            "w",
            compression=zipfile.ZIP_DEFLATED,
            compresslevel=9,
        ) as target:
            for name in sorted(source.namelist()):
                original = source.getinfo(name)
                info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
                info.compress_type = zipfile.ZIP_DEFLATED
                info.external_attr = original.external_attr
                info.create_system = 0
                target.writestr(info, source.read(name), compress_type=zipfile.ZIP_DEFLATED, compresslevel=9)


def _write_hashes(paths: tuple[Path, ...]) -> None:
    lines = []
    for path in paths:
        digest = hashlib.sha256(path.read_bytes()).hexdigest()
        lines.append(f"{digest}  {path.name}")
    (PACKAGE_DIR / "SHA256SUMS.txt").write_text("\n".join(lines) + "\n", encoding="utf-8", newline="\n")


def build_all() -> tuple[Path, ...]:
    builders = (
        (build_main_manuscript, MAIN_DOCX),
        (build_title_page, TITLE_DOCX),
        (build_cover_letter, COVER_DOCX),
        (build_supplement, SUPPLEMENT_DOCX),
    )
    for builder, destination in builders:
        _deterministic_save(builder(), destination)
    _write_hashes(OUTPUTS)
    return OUTPUTS


if __name__ == "__main__":
    for output in build_all():
        print(output.relative_to(ROOT))
