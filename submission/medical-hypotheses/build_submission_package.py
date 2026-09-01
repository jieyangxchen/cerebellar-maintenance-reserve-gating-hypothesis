#!/usr/bin/env python3
"""Build the deterministic Word package for Medical Hypotheses."""

from __future__ import annotations

import importlib.util
import re
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
PACKAGE_DIR = Path(__file__).resolve().parent / "package"
LEGACY_BUILDER = ROOT / "submission" / "the-cerebellum" / "build_submission_package.py"
PRIVACY_SCRUBBER = (
    Path.home()
    / ".codex/plugins/cache/openai-primary-runtime/documents/26.826.12353/skills"
    / "documents/scripts/privacy_scrub.py"
)

spec = importlib.util.spec_from_file_location("legacy_submission_builder", LEGACY_BUILDER)
if spec is None or spec.loader is None:
    raise RuntimeError(f"Unable to load shared document helpers from {LEGACY_BUILDER}")
base = importlib.util.module_from_spec(spec)
spec.loader.exec_module(base)

privacy_spec = importlib.util.spec_from_file_location("docx_privacy_scrubber", PRIVACY_SCRUBBER)
if privacy_spec is None or privacy_spec.loader is None:
    raise RuntimeError(f"Unable to load DOCX privacy scrubber from {PRIVACY_SCRUBBER}")
privacy = importlib.util.module_from_spec(privacy_spec)
privacy_spec.loader.exec_module(privacy)


TITLE = "The Maintenance–Reserve–Gating Hypothesis of Hereditary Cerebellar Ataxia"
SUBTITLE = "A testable model of non-linear modifier effects and delayed phenoconversion"
RUNNING_TITLE = "Maintenance–reserve–gating hypothesis in hereditary ataxia"
VERSION = "0.3.1"
PACKAGE_DATE = "1 September 2026"
EVIDENCE_CUTOFF = "19 August 2026"
FIXED_TIME = datetime(2026, 9, 1, 0, 0, 0, tzinfo=timezone.utc)

HIGHLIGHTS_SOURCE = Path(__file__).resolve().parent / "highlights.txt"
LEGACY_MAIN_DOCX = PACKAGE_DIR / "01_Main_Manuscript.docx"
MAIN_DOCX = PACKAGE_DIR / "01_Anonymized_Manuscript.docx"
TITLE_DOCX = PACKAGE_DIR / "02_Title_Page.docx"
COVER_DOCX = PACKAGE_DIR / "03_Cover_Letter.docx"
HIGHLIGHTS_DOCX = PACKAGE_DIR / "04_Highlights.docx"
CREDIT_DOCX = PACKAGE_DIR / "05_CRediT_Author_Statement.docx"
INTEREST_DOCX = PACKAGE_DIR / "06_Declaration_of_Interest.docx"
ETHICS_DOCX = PACKAGE_DIR / "07_Ethics_Statement.docx"
OUTPUTS = (
    MAIN_DOCX,
    TITLE_DOCX,
    COVER_DOCX,
    HIGHLIGHTS_DOCX,
    CREDIT_DOCX,
    INTEREST_DOCX,
    ETHICS_DOCX,
)

ETHICS_STATEMENT = (
    "Ethics approval and informed consent were not required because this article "
    "is a hypothesis and critical review that reports no original research involving "
    "human participants, animals, participant-level data, biological material, or "
    "identifiable personal information."
)
CREDIT_STATEMENT = (
    "Jieyang Chen: Conceptualization; Methodology; Investigation; Data curation; "
    "Software; Visualization; Writing – original draft; Writing – review & editing; "
    "Project administration."
)
INTEREST_STATEMENT = "Declarations of interest: none."
REPOSITORY_LICENSING_STATEMENT = (
    "For transparency, the author retains copyright in the associated public "
    "repository and may consider future commercial-licensing requests. This is a "
    "licensing position rather than a current financial or personal interest; no "
    "commercial funding or payment was received for this work."
)
AI_STATEMENT = (
    "During preparation, the author used OpenAI Codex for manuscript organization, "
    "language and readability editing, figure and repository-check code, and internal "
    "consistency review. The author reviewed and edited the output and remains fully "
    "responsible for the submitted work. OpenAI Codex is not an author."
)


def _read_highlights() -> tuple[str, ...]:
    highlights = tuple(
        line[2:].strip()
        for line in HIGHLIGHTS_SOURCE.read_text(encoding="utf-8").splitlines()
        if line.startswith("- ")
    )
    if not 3 <= len(highlights) <= 5:
        raise ValueError("Highlights must contain 3 to 5 bullet points")
    too_long = [item for item in highlights if len(item) > 85]
    if too_long:
        raise ValueError(f"Highlights exceed 85 characters: {too_long}")
    return highlights


def _anonymize_body(body: str) -> str:
    body = body.replace(
        "The public repository identified on the title page",
        "A supporting repository disclosed separately to the editorial office",
    )
    body = body.replace("The repository's trial concepts", "The supporting package's trial concepts")
    body = body.replace("Trial concepts in this repository", "Trial concepts in the supporting package")
    anonymous_data_statement = (
        "### Data and code availability\n\n"
        "No participant-level or experimental data are associated with this article. "
        "Figure source, deterministic sample-size scenarios, protocols, the evidence "
        "audit, and repository checks are deposited in an author-identifying public "
        "repository disclosed separately to the editorial office. The direct URL is "
        "omitted from this reviewer manuscript to preserve double-anonymized review. "
        "The research-planning documents are not substitutes for sponsor-approved "
        "protocols or statistical analysis plans.\n\n"
    )
    body = re.sub(
        r"(?ms)^### Data and code availability\s+.*?(?=^### Ethics statement)",
        anonymous_data_statement,
        body,
    )
    body = re.sub(
        r"(?ms)^### CRediT author statement\s+.*?(?=^### )",
        "",
        body,
    )
    body = re.sub(
        r"(?ms)^### Acknowledgements\s+.*?(?=^### )",
        "",
        body,
    )
    return body.strip()


def _extract_main_sections(text: str) -> tuple[str, str, str]:
    abstract_marker = "## Abstract"
    body_marker = "## Hypothesis and evidentiary boundary"
    abstract_start = text.index(abstract_marker) + len(abstract_marker)
    body_start = text.index(body_marker)
    abstract_and_keywords = text[abstract_start:body_start].strip()
    keyword_match = re.search(r"\*\*Keywords:\*\*\s*(.+)$", abstract_and_keywords, re.MULTILINE)
    if keyword_match is None:
        raise ValueError("Keywords not found in manuscript")
    keywords = keyword_match.group(1).strip()
    abstract = abstract_and_keywords[: keyword_match.start()].strip()
    body = _anonymize_body(text[body_start:].strip())
    return abstract, keywords, body


def _add_anonymous_main_title_page(doc: Document) -> None:
    paragraph = doc.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base._add_inline_runs(paragraph, TITLE, size=14, default_bold=True)
    base._add_centered(doc, SUBTITLE, size=11, italic=True, after=10)
    base._add_centered(doc, "Article type: Hypothesis", size=9.5, after=3)
    base._add_centered(doc, f"Running title: {RUNNING_TITLE}", size=9.5, italic=True, after=3)
    base._add_centered(
        doc,
        f"Manuscript version {VERSION}; evidence cut-off {EVIDENCE_CUTOFF}",
        size=9.5,
        after=3,
    )
    base._add_centered(
        doc,
        "Anonymous manuscript for double-anonymized peer review. This hypothesis article reports no original human-participant or animal research.",
        size=9.5,
        italic=True,
        after=8,
    )
    doc.add_page_break()


def _configure_shared_globals() -> None:
    base.PACKAGE_DIR = PACKAGE_DIR
    base.TITLE = TITLE
    base.RUNNING_TITLE = RUNNING_TITLE
    base.VERSION = VERSION
    base.PACKAGE_DATE = PACKAGE_DATE
    base.EVIDENCE_CUTOFF = EVIDENCE_CUTOFF
    base.FIXED_TIME = FIXED_TIME
    base.MAIN_DOCX = MAIN_DOCX
    base.TITLE_DOCX = TITLE_DOCX
    base.COVER_DOCX = COVER_DOCX
    base.OUTPUTS = OUTPUTS
    base._extract_main_sections = _extract_main_sections
    base._add_main_title_page = _add_anonymous_main_title_page


def _normalize_footer(doc: Document) -> None:
    """Use a centered numeric field so every rendered page number stays visible."""

    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.right_indent = Inches(0)
            for node in paragraph._p.iter():
                if node.tag.endswith("}t") and node.text == "Page ":
                    node.text = ""


def build_main_manuscript() -> Document:
    doc = base.build_main_manuscript()
    doc.core_properties.subject = "Anonymous hypothesis article for Medical Hypotheses"
    _normalize_footer(doc)
    return doc


def build_title_page() -> Document:
    doc = Document()
    base._configure_document(
        doc,
        title=f"Title page - {TITLE}",
        subject="Separate submission title page",
        header_text="Title page",
    )
    paragraph = doc.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base._add_inline_runs(paragraph, TITLE, size=14, default_bold=True)
    base._add_centered(doc, SUBTITLE, size=11, italic=True, after=10)
    base._add_centered(doc, base.AUTHOR, size=11, bold=True, after=3)
    base._add_centered(doc, f"ORCID: {base.ORCID}", size=10, after=2)
    base._add_centered(doc, base.AFFILIATION, size=10, after=2)
    base._add_centered(
        doc,
        f"Corresponding author: {base.AUTHOR} | {base.EMAIL}",
        size=10,
        after=12,
    )

    entries = (
        ("Running title", RUNNING_TITLE),
        ("Proposed article type", "Hypothesis"),
        ("Research status", "Hypothesis and critical review; no original participant-level or animal research"),
        ("Peer-review file", "A separate anonymized manuscript with scrubbed Word metadata is supplied for double-anonymized review"),
        ("Manuscript version", VERSION),
        ("Evidence cut-off", EVIDENCE_CUTOFF),
        ("Public preprint and repository", base.REPOSITORY),
        ("Funding", "This work received no external funding."),
        (
            "Declaration of interest",
            INTEREST_STATEMENT,
        ),
        ("Repository licensing", REPOSITORY_LICENSING_STATEMENT),
        (
            "CRediT author statement",
            CREDIT_STATEMENT,
        ),
        ("Ethics statement", ETHICS_STATEMENT),
        ("Acknowledgements", "None."),
        ("AI-assisted tools", AI_STATEMENT),
    )
    for label_text, value in entries:
        paragraph = doc.add_paragraph()
        label = paragraph.add_run(f"{label_text}: ")
        base._set_run_font(label, size=10, bold=True)
        base._add_inline_runs(paragraph, value, size=10)
    _normalize_footer(doc)
    return doc


def build_cover_letter() -> Document:
    doc = Document()
    base._configure_document(
        doc,
        title=f"Cover letter - {TITLE}",
        subject="Cover letter to Medical Hypotheses",
        header_text="Cover letter | Medical Hypotheses",
    )

    for line in (PACKAGE_DATE, "Editor-in-Chief and Editors", "Medical Hypotheses"):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        base._add_inline_runs(paragraph, line, size=10)
    doc.add_paragraph()
    salutation = doc.add_paragraph()
    base._add_inline_runs(salutation, "Dear Editor-in-Chief and Editors,", size=10)

    paragraphs = (
        f"I am submitting the manuscript entitled “{TITLE}” for consideration as a hypothesis article in Medical Hypotheses.",
        "The article proposes a nested, falsifiable explanation for timing heterogeneity in hereditary cerebellar ataxia. Its directly testable layer asks whether one outcome-blind, prospectively frozen measured exposure score shows a reproducible non-linear and genotype-specific association with phenoconversion and biomarker trajectories. Its optional mechanistic extension posits an unidentified input, activation/exchange gating, and dynamic reserve. The manuscript states explicitly that a statistical hump alone would not demonstrate gating.",
        "The hypothesis builds on established structural and functional concepts of cerebellar reserve. Its narrower novelty is the proposed activation–supply mismatch, the observation model linking dynamic reserve to phenoconversion, and layer-specific rejection rules. SCA3 is the primary proposed test bed and SCA6 a stringent transport test. The full protocol and statistical-analysis concepts remain in the public repository rather than being presented as validated methods or implementation-ready trials.",
        "The article reports no participant-level data and no new human or animal experiments. Current evidence supports component premises but does not demonstrate the unknown input, a gate, outward loss, or a hump-shaped human exposure curve. Those limitations, competing explanations, and results that would reject each layer are stated in the abstract and main text. The manuscript is a research hypothesis, not a treatment recommendation.",
        "The submission package follows double-anonymized review requirements: the reviewer manuscript omits author identity, affiliation, contact details, ORCID, acknowledgements, named CRediT attribution, and the identifying repository URL. Its Word metadata has also been scrubbed. The title page, Highlights, CRediT statement, Declaration of Interest, and Ethics statement are supplied separately.",
        f"A preprint and supporting research-planning package are publicly available under a non-commercial licence at {base.REPOSITORY}. This history is disclosed transparently. The version submitted here is not under consideration by another journal.",
        "No external funding was received. The author declares no competing interests and no acknowledgements. The manuscript includes a declaration of the use of OpenAI Codex during preparation; the named author has reviewed the complete article and takes full responsibility for all scientific claims, references, figures, and wording.",
        "Thank you for considering this hypothesis.",
    )
    for text in paragraphs:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        base._add_inline_runs(paragraph, text, size=10)

    signature = doc.add_paragraph()
    signature.paragraph_format.space_before = Pt(5)
    signature.paragraph_format.space_after = Pt(0)
    signature.paragraph_format.keep_together = True
    for index, line in enumerate(
        ("Sincerely,", base.AUTHOR, base.AFFILIATION, base.EMAIL, f"ORCID: {base.ORCID}")
    ):
        if index:
            signature.add_run().add_break()
        base._add_inline_runs(signature, line, size=10)
    _normalize_footer(doc)
    return doc


def _build_statement_document(
    heading_text: str,
    subject: str,
    paragraphs: tuple[str, ...],
    *,
    bullets: tuple[str, ...] = (),
) -> Document:
    doc = Document()
    base._configure_document(
        doc,
        title=f"{heading_text} - {TITLE}",
        subject=subject,
        header_text=heading_text,
    )
    heading = doc.add_paragraph(style="Title")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base._add_inline_runs(heading, heading_text, size=14, default_bold=True)
    base._add_centered(doc, TITLE, size=10.5, italic=True, after=12)
    for text in paragraphs:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        base._add_inline_runs(paragraph, text, size=10)
    for text in bullets:
        paragraph = doc.add_paragraph(style="List Bullet")
        base._add_inline_runs(paragraph, text, size=10)
    _normalize_footer(doc)
    return doc


def build_highlights() -> Document:
    return _build_statement_document(
        "Highlights",
        "Highlights for Medical Hypotheses",
        (),
        bullets=_read_highlights(),
    )


def build_credit_statement() -> Document:
    return _build_statement_document(
        "CRediT author statement",
        "Contributor role statement",
        (CREDIT_STATEMENT,),
    )


def build_interest_statement() -> Document:
    return _build_statement_document(
        "Declaration of Interest",
        "Declaration of interest",
        (INTEREST_STATEMENT, REPOSITORY_LICENSING_STATEMENT),
    )


def build_ethics_statement() -> Document:
    return _build_statement_document(
        "Ethics statement",
        "Ethics statement",
        (ETHICS_STATEMENT,),
    )


def _canonicalize_existing_docx(source: Path, destination: Path) -> None:
    with zipfile.ZipFile(source, "r") as archive, zipfile.ZipFile(
        destination,
        "w",
        compression=zipfile.ZIP_DEFLATED,
        compresslevel=9,
    ) as target:
        for name in sorted(archive.namelist()):
            original = archive.getinfo(name)
            info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = original.external_attr
            info.create_system = 0
            target.writestr(
                info,
                archive.read(name),
                compress_type=zipfile.ZIP_DEFLATED,
                compresslevel=9,
            )


def _scrub_word_metadata(destination: Path) -> None:
    with tempfile.TemporaryDirectory(prefix="medical-hypotheses-privacy-") as temporary:
        scrubbed = Path(temporary) / destination.name
        privacy.scrub(str(destination), str(scrubbed))
        _canonicalize_existing_docx(scrubbed, destination)


def build_all() -> tuple[Path, ...]:
    _configure_shared_globals()
    LEGACY_MAIN_DOCX.unlink(missing_ok=True)
    builders = (
        (build_main_manuscript, MAIN_DOCX),
        (build_title_page, TITLE_DOCX),
        (build_cover_letter, COVER_DOCX),
        (build_highlights, HIGHLIGHTS_DOCX),
        (build_credit_statement, CREDIT_DOCX),
        (build_interest_statement, INTEREST_DOCX),
        (build_ethics_statement, ETHICS_DOCX),
    )
    for builder, destination in builders:
        base._deterministic_save(builder(), destination)
        _scrub_word_metadata(destination)
    base._write_hashes(OUTPUTS)
    return OUTPUTS


if __name__ == "__main__":
    for output in build_all():
        print(output.relative_to(ROOT))
